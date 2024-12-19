from PyQt6.QtWidgets import QFileDialog,QApplication, QMainWindow, QTableView, QWidget, QPushButton, QLineEdit, QMessageBox, QDateEdit, QDialog, QComboBox, QMenuBar, QMenu, QTableWidgetItem, QTableView, QHeaderView, QVBoxLayout, QLabel
from PyQt6.QtCore import QDateTime, Qt, QDate, QEvent
from PyQt6.QtGui import QAction, QStandardItemModel, QStandardItem

from datetime import date, datetime
import traceback
from dis_form.mainW import Ui_MainWindow

from class_database import Database
from class_wid_add_update import widget_add
from class_widget_quntity import widget_add_kol_vo
from class_image import ImageWindow
import os
from openpyxl import Workbook
from docx import Document

from dis_form.dis_login import Ui_Form
from dis_form.dis_wid_add_order import Ui_Dialog_order
from dis_form.dis_wid_add_report import Ui_Dialog_report
from class_database import Database

class LoginForm(QWidget, Ui_Form):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.datab = Database()
        self.but_login.clicked.connect(self.entry)

    def entry(self):
        data = self.datab.query_entry(self.r_le_login.text(), self.r_le_password.text(),
                                      self.r_comboBox_position.currentText())
        print(data, self.r_le_login.text(), self.r_le_password.text(), self.r_comboBox_position.currentText())
        if not data:
            QMessageBox.critical(self, 'Ошибка', 'Неправильные данные!')
        else:
            self.openMainWindow(self.r_comboBox_position.currentText(), data[0])

    def openMainWindow(self, position, id_users):
        self.mainWindow = MainWindow(position, id_users)
        self.mainWindow.show()
        if position == 'Сотрудник':
            self.mainWindow.action_workers.setVisible(False)
        self.close()

class MainWindow(QMainWindow, Ui_MainWindow):
    def __init__(self, position, id_users):
        super().__init__()
        self.setupUi(self)
        self.setWindowTitle("Главное окно")
        self.showMaximized()
        self.tableView.verticalHeader().setVisible(False)


        self.datab = Database()
        self.table_name = 'Автомобили'
        self.position = position
        self.id_users = id_users

        self.wid_add_or_update = QDialog()
        self.widget_kol = QDialog()

        self.values_sales_delivers = []
        self.sales_or_delivers = None
        self.sum = 0



        self.tableView.clicked.connect(self.opr_click_tableview)


        self.but_search.clicked.connect(self.search_column)
        self.but_excel.clicked.connect(self.export_table_to_excel_with_header)
        self.but_delete.clicked.connect(lambda : self.delete_column())
        self.but_add.clicked.connect(lambda :self.add_or_update_line())
        self.but_redac.clicked.connect(lambda :self.add_or_update_line())
        self.but_profile.clicked.connect(lambda :self.add_or_update_line())
        self.but_order.clicked.connect(lambda :self.open_widget_add_order())
        self.but_dogovor.clicked.connect(lambda: self.create_dogovor())
        self.but_report.clicked.connect(lambda: self.open_widget_report())




        self.action_about.triggered.connect(self.show_about_info)
        self.action_exit.triggered.connect(self.close_all_windows)

        table_names = self.datab.get_table_name()
        self.tables = []
        if self.position == 'Сотрудник':
            table_names.remove('Сотрудники')
            self.but_report.setVisible(False)
        for name in table_names:
            action = QAction(name, self)
            self.tables.append(action)
            self.m_tables.addAction(action)
            action.triggered.connect(lambda checked, table_name=name: self.opr_table(table_name))

        self.opr_table(self.table_name)


    def open_widget_report(self):
        try:
            self.wid_report = widget_report()
            values = ['Выручка за период', 'Прибыль за период', 'Продажи сотрудников за период']

            for value in values:
                self.wid_report.cb_name_report.addItem(value)
            self.wid_report.but_save_add.clicked.connect(lambda: self.create_report())

            self.wid_report.show()

        except Exception as e:
            QMessageBox.information(self, 'Внимание', 'Заполните поле для поиска!')
            print(f"2 Ошибка при выполнении запроса: {e}")
            traceback.print_exc()

    def create_report(self):
        try:
            name_report = self.wid_report.cb_name_report.currentText()

            # Получаем даты из виджетов и преобразуем в строковый формат
            date = self.wid_report.date_start.date()
            date_start = date.toString(Qt.DateFormat.ISODate)

            date = self.wid_report.date_end.date()
            date_end = date.toString(Qt.DateFormat.ISODate)

            # Формируем SQL запрос в зависимости от выбранного отчета
            if name_report == 'Выручка за период':
                view_name = 'Отчет_Выручка_За_Период'

            elif name_report == 'Прибыль за период':
                view_name = 'Отчет_Прибыль_За_Период'

            elif name_report == 'Продажи сотрудников за период':
                view_name = 'Отчет_Продажи_Сотрудников'

            else:
                QMessageBox.information(self, 'Уведомление', 'Отчет не найден!')
                return

            # Отправляем запрос на вывод в таблицу
            self.wid_report.close()
            self.export_report_to_excel(view_name, date_start, date_end)  # Выводим данные таблицы

        except Exception as e:
            print(f"22 Ошибка при выполнении запроса: {e}")


    def export_report_to_excel(self, view_name, date_start, date_end):
        try:
            directory = 'excel'
            filename = f'{view_name} с {date_start} по {date_end}.xlsx'
            file = os.path.join(directory, filename)

            if not os.path.exists(directory):
                os.makedirs(directory)  # Создаем директорию, если её нет

            data, columns = self.datab.get_data_table_with_date(view_name, date_start, date_end)

            # Создание Excel-книги
            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Exported Data"

            # Запись заголовков в первую строку
            sheet.append(columns)

            # Запись данных в Excel
            for row in data:
                sheet.append(row)

            workbook.save(file)  # Сохраняем Excel-файл
            QMessageBox.information(self, 'Внимание', f'Отчет успешно создан!')

        except Exception as e:
            print(f"222 Ошибка при выполнении запроса: {e}")


    def create_dogovor(self):
        try:
            id_order, _, column = self.opr_cell_row_column(self.tableView)
            if not id_order or column != 0:
                QMessageBox.information(self, 'Внимание', 'Выберите продажу!')
                return

            # Получение данных о заказе
            data_order = self.datab.query_table('Продажи', id_order)

            id_client = int(data_order[1])
            sql_query = f'''
            SELECT 
                Имя,
                Фамилия,
                Телефон,
                Email,
                Адрес,
                Серия_паспорта,
                Номер_паспорта,
                Дата_выдачи,
                Кем_выдан
            FROM Клиенты
            WHERE ID={id_client}'''

            self.datab.cursor.execute(sql_query)
            data_clients = self.datab.cursor.fetchone()

            id_manager = int(data_order[2])
            sql_query = f'''
            SELECT 
                Имя,
                Фамилия,
                Телефон,
                Email,
                Должность
            FROM Сотрудники
            WHERE ID={id_manager}'''

            self.datab.cursor.execute(sql_query)
            data_manager = self.datab.cursor.fetchone()

            id_car = data_order[3]
            sql_query = f'''
            SELECT 
                Тип_марки.Название AS Марка,
                Тип_модели.Название AS Модель,
                Автомобили.Год,
                Автомобили.VIN,
                Автомобили.Цвет,
                Тип_двигателя.Название AS Тип_двигателя,
                Автомобили.Объем,
                Автомобили.Мощность,
                Тип_трансмиссии.Название AS Трансмиссия,
                Тип_привода.Название AS Привод,
                Автомобили.Цена
            FROM Автомобили
            LEFT JOIN Тип_модели ON Автомобили.ID_Тип_модели = Тип_модели.ID
            LEFT JOIN Тип_марки ON Тип_модели.ID_Тип_марки = Тип_марки.ID
            LEFT JOIN Тип_двигателя ON Автомобили.ID_Тип_двигателя = Тип_двигателя.ID
            LEFT JOIN Тип_трансмиссии ON Автомобили.ID_Тип_трансмиссии = Тип_трансмиссии.ID
            LEFT JOIN Тип_привода ON Автомобили.ID_Тип_привода = Тип_привода.ID
            WHERE Автомобили.ID={id_car}'''

            self.datab.cursor.execute(sql_query)
            data_car = self.datab.cursor.fetchone()

            id_type_pays = data_order[6]
            sql_query = f'''
            SELECT 
                Название AS Тип_оплаты
            FROM Тип_оплаты
            WHERE ID={id_type_pays}'''

            self.datab.cursor.execute(sql_query)
            data_type_pays = self.datab.cursor.fetchone()

            datatime_order = data_order[4]
            price = data_order[5]

            # Данные компании
            data_company = ['Автомир', 'г.Москва, пр.Победы, д.20', '+74951234567']

            # Создание документа
            doc = Document()

            # Заголовок
            title = doc.add_heading(level=1)
            title_run = title.add_run('ДОГОВОР КУПЛИ-ПРОДАЖИ АВТОМОБИЛЯ')


            # Ввод основной информации
            doc.add_paragraph(f'г. Москва, дата: {datatime_order}')
            doc.add_paragraph(f'\nПродавец: {data_company[0]}, адрес: {data_company[1]}, телефон: {data_company[2]}.')
            doc.add_paragraph(
                f'Покупатель: {data_clients[0]} {data_clients[1]}, телефон: {data_clients[2]}, паспорт: {data_clients[5]} {data_clients[6]}, выдан {data_clients[7]} {data_clients[8]}.'
            )

            doc.add_paragraph(
                f'Продавец обязуется передать, а Покупатель принять и оплатить следующий автомобиль:\n'
                f'Марка: {data_car[0]}, Модель: {data_car[1]}, Год: {data_car[2]}, VIN: {data_car[3]}, Цвет: {data_car[4]},\n'
                f'Тип двигателя: {data_car[5]}, Объем: {data_car[6]}, Мощность: {data_car[7]} л.с., Трансмиссия: {data_car[8]}, Привод: {data_car[9]}.\n'
            )

            doc.add_paragraph(f'Стоимость автомобиля: {price} руб.')

            doc.add_paragraph(
                f'''Условия оплаты:
    1. Покупатель обязуется оплатить полную стоимость автомобиля до момента передачи.
    2. Оплата производится путем {data_type_pays[0]}.

Права и обязанности сторон:
    1. Право собственности на транспортное средство переходит к Покупателю с момента подписания настоящего договора.
    2. Покупатель обязуется принять автомобиль в состоянии, указанном в акте приема-передачи.\n\nСотрудник (Менеджер): {data_manager[0]} {data_manager[1]}.\n''')

            doc.add_paragraph('Настоящий договор составлен в двух экземплярах, имеющих равную юридическую силу.')

            # Подписи сторон
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = 'Продавец: ________________'
            table.cell(0, 1).text = 'Покупатель: ________________'
            table.cell(1, 0).text = f'({data_company[0]})'
            table.cell(1, 1).text = f'({data_clients[0]} {data_clients[1]})'

            # Сохранение документа
            file_name = f'word/Договор_{id_order}.docx'
            doc.save(file_name)


            QMessageBox.information(self, 'Успех', f'Договор успешно создан.')

        except Exception as e:
            QMessageBox.critical(self, 'Ошибка', f'Не удалось создать договор: {str(e)}')



        except Exception as e:
            QMessageBox.information(self, 'Внимание', 'Заполните поле для поиска!')
            print(f"2 Ошибка при выполнении запроса: {e}")
            traceback.print_exc()

    def open_widget_add_order(self):
        try:
            id_car, _, column = self.opr_cell_row_column(self.tableView)
            if not id_car or column != 0:
                QMessageBox.information(self, 'Внимание', 'Выберите автомобиль для оформления!')
                return

            self.wid_add_order = widget_add_order()

            sql_query_client = '''SELECT Клиенты.ID, Клиенты.Имя, Клиенты.Фамилия, Клиенты.Серия_паспорта, Клиенты.Номер_паспорта FROM Клиенты;'''
            sql_query_type_pay = '''SELECT Тип_оплаты.ID, Тип_оплаты.Название FROM Тип_оплаты;'''

            # Получение данных о клиентах
            self.datab.cursor.execute(sql_query_client)
            data_clients = self.datab.cursor.fetchall()
            for client in data_clients:
                client_id, first_name, last_name, passport_series, passport_number = client
                formatted_client = f"{first_name} {last_name} (Паспорт: {passport_series} {passport_number})"
                self.wid_add_order.cb_clients.addItem(formatted_client,
                                                      userData=client_id)  # Добавляем userData для хранения ID

            # Получение данных о типах оплаты
            self.datab.cursor.execute(sql_query_type_pay)
            data_type_pays = self.datab.cursor.fetchall()
            for payment in data_type_pays:
                payment_id, payment_name = payment
                self.wid_add_order.cb_pay.addItem(payment_name, userData=payment_id)  # userData для ID

            # Установка цены
            self.wid_add_order.le_price.setText(f"{self.datab.get_price(self.table_name, id_car)}")

            print(data_clients)
            print(data_type_pays)


            self.wid_add_order.but_save_add.clicked.connect(lambda: self.save_order(id_car, self.id_users))
            self.wid_add_order.show()

        except Exception as e:
            QMessageBox.information(self, 'Внимание', 'Заполните поле для поиска!')
            print(f"2 Ошибка при выполнении запроса: {e}")
            traceback.print_exc()

    def save_order(self, id_car, id_users):
        try:
            table_name = 'Продажи'
            client_id = self.wid_add_order.cb_clients.currentData()
            payment_id = self.wid_add_order.cb_pay.currentData()
            try:
                price = float(self.wid_add_order.le_price.text())
            except ValueError:
                self.wid_add_order.le_price.setText('')
                self.wid_add_order.le_price.setPlaceholderText('Неправильные данные!')
                return

            id = self.datab.incr(table_name)
            current_datetime = datetime.now()
            formatted_datetime = current_datetime.strftime("%Y-%m-%d %H:%M:%S")

            print(client_id, id_users, id_car, price, payment_id)
            values = [id, client_id, id_users, id_car, formatted_datetime, price, payment_id]
            self.datab.query_add_new_column(table_name, values)
            self.wid_add_order.close()
            QMessageBox.information(self, 'Внимание', 'Операция прошла успешно!')
            self.opr_sql_query()

        except Exception as e:
            print(f"1 Ошибка при выполнении запроса: {e}")  # Обработка ошибок

    def print_structure_sales_deliveries_in_word(self):
        try:
            id_sales_delivers = self.values_sales_delivers[0]
            table_name = self.sales_or_delivers[0]
            structure_table_name = self.sales_or_delivers[1]
            column = self.sales_or_delivers[2]
            # Получаем дату и время для отчета
            if id_sales_delivers:
                self.datab.cursor.execute(f'SELECT date FROM {table_name} WHERE id = {id_sales_delivers}')
                time = self.datab.cursor.fetchone()[0]
            else:
                time = QDateTime.currentDateTime().toString('dd.MM.yy hh-mm-ss')  # Заменяем ":" на "-"

            if table_name == 'sales':
                sql_query = f'''SELECT sweets.id, sweets.name, sweets.price, sale_items.quantity, unit.name as unit_name, manufacturers.name as manufacturers_name
                                FROM sale_items JOIN sweets ON sale_items.sweets_id = sweets.id 
                                JOIN unit on unit.id = sweets.unit_id 
                                JOIN manufacturers ON manufacturers.id = sweets.manufacturers_id
                                WHERE sale_items.sales_id = {id_sales_delivers};'''
            else:
                sql_query = f'''SELECT sweets.id, sweets.name, sweets.price, delivery_items.quantity, unit.name as unit_name, manufacturers.name as manufacturers_name
                                FROM delivery_items JOIN sweets ON delivery_items.sweets_id = sweets.id 
                                JOIN unit on unit.id = sweets.unit_id 
                                JOIN manufacturers ON manufacturers.id = sweets.manufacturers_id
                                WHERE delivery_items.deliveries_id = {id_sales_delivers};'''

            self.datab.cursor.execute(sql_query)
            data = self.datab.cursor.fetchall()

            doc = Document()
            report = ['Продажа', 'Покупатель', 'Продавец'] if table_name == 'sales' else ['Поставка', 'Принял',
                                                                                          'Поставщик']
            doc.add_heading(f'{report[0]} №{id_sales_delivers} - {time}\n', level=1).paragraph_format.alignment = 1
            # Добавляем заголовки столбцов
            columns = [description[0] for description in self.datab.cursor.description]
            table = doc.add_table(rows=1, cols=len(columns))
            table.style = 'TableGrid'
            for col_num, column in enumerate(columns):
                table.cell(0, col_num).text = column
            # Добавляем данные в таблицу
            for row in data:
                row_cells = table.add_row().cells
                for col_num, value in enumerate(row):
                    row_cells[col_num].text = str(value)
            # Добавляем общую сумму
            total_amount_paragraph = doc.add_paragraph()
            total_amount_run = total_amount_paragraph.add_run('Общая сумма:')
            total_amount_paragraph.add_run(f' {self.sum}\n').bold = True
            total_amount_paragraph.alignment = 2

            buyer = doc.add_paragraph()
            buyer = buyer.add_run(f'{report[1]}: __________________  ________')
            buyer.alignment = 0
            seller = doc.add_paragraph()
            seller = seller.add_run(f'{report[2]}: __________________  ________')
            seller.alignment = 0
            doc.save(f'word/{report[0]} №{id_sales_delivers}.docx')

            self.widget_add_ord.message_word()
        except Exception as e:
            print(f"1 Ошибка при выполнении запроса: {e}")  # Обработка ошибок

    def add_or_update_line(self):
        try:
            table_name = self.table_name
            sender = self.sender()  # Получаем отправителя сигнала
            function_name = sender.text()  # Получаем название функции (кнопки)
            id, _, column = self.opr_cell_row_column(self.tableView)  # Получаем id, колонку и другие данные

            if function_name == '🖋️ Изменить' and (not id or column != 0):
                QMessageBox.information(self, 'Внимание', 'Выберите строку для изменения!')
                return

            if function_name == '👤':
                id = self.id_users
                function_name = '🖋️ Изменить'
                table_name = 'Сотрудники'

            self.wid_add_or_update = widget_add(table_name, function_name, self.position, self)

            self.wid_add_or_update.open_widget_add(id)

        except Exception as e:
            print(f"10 Ошибка при выполнении запроса: {e}")


    def opr_table(self, table_name):
        self.table_name = table_name
        if self.table_name in ['Продажи']:
            self.but_add.setVisible(False)
            self.but_redac.setVisible(False)
            self.but_order.setVisible(False)
            self.but_dogovor.setVisible(True)
            if self.position == 'Руководитель':
                self.but_delete.setVisible(True)
        else:
            self.but_add.setVisible(True)
            self.but_redac.setVisible(True)
            self.but_dogovor.setVisible(False)
            self.but_delete.setVisible(False)
            self.but_order.setVisible(False)
            if self.table_name == 'Автомобили':
                self.but_order.setVisible(True)

        self.opr_sql_query()

    def opr_sql_query(self):
        if self.table_name == 'Автомобили':
            sql_query = f'''SELECT 
                Автомобили.ID AS ID_Автомобиля,
                Тип_марки.Название AS Марка,
                Автомобили.ID_Тип_модели AS ID_Модели,
                Тип_модели.Название AS Модель,
                Автомобили.Год AS Год,
                Автомобили.VIN AS VIN,
                Автомобили.Цвет AS Цвет,
                Автомобили.Фото AS Фото,
                Автомобили.ID_Тип_двигателя AS ID_Двигателя,
                Тип_двигателя.Название AS Тип_двигателя,
                Автомобили.Объем AS Объем_двигателя,
                Автомобили.Мощность AS Мощность_двигателя,
                Автомобили.ID_Тип_трансмиссии AS ID_Трансмиссии,
                Тип_трансмиссии.Название AS Тип_трансмиссии,
                Автомобили.ID_Тип_привода AS ID_Привода,
                Тип_привода.Название AS Тип_привода,
                Автомобили.Масса AS Масса,
                Автомобили.Закуп AS Закупочная_цена,
                Автомобили.Цена AS Продажная_цена,
                Автомобили.Комплектация AS Комплектация,
                Автомобили.Статус AS Статус
            FROM 
                Автомобили
            JOIN 
                Тип_модели ON Автомобили.ID_Тип_модели = Тип_модели.ID
            JOIN 
                Тип_марки ON Тип_модели.ID_Тип_марки = Тип_марки.ID
            JOIN 
                Тип_двигателя ON Автомобили.ID_Тип_двигателя = Тип_двигателя.ID
            JOIN 
                Тип_трансмиссии ON Автомобили.ID_Тип_трансмиссии = Тип_трансмиссии.ID
            JOIN 
                Тип_привода ON Автомобили.ID_Тип_привода = Тип_привода.ID Where Автомобили.Статус = 'Активный';
            '''

        elif self.table_name == 'Продажи' and self.position == 'Сотрудник':
            sql_query = f'''SELECT 
                            Продажи.ID AS ID_Продажи,
                            Клиенты.ID AS ID_Клиента,
                            Клиенты.Имя || ' ' || Клиенты.Фамилия AS Клиент,    
                            Сотрудники.ID AS ID_Сотрудника,  
                            Сотрудники.Имя || ' ' || Сотрудники.Фамилия AS Сотрудник,
                            Автомобили.ID AS ID_Автомобиля,
                            Автомобили.VIN AS VIN_Автомобиля,
                            Тип_модели.Название AS Модель_Автомобиля,
                            Тип_марки.Название AS Марка_Автомобиля,
                            Продажи.Дата_продажи AS Дата_Продажи,
                            Продажи.Цена AS Цена_Продажи,
                            Тип_оплаты.Название AS Тип_Оплаты
                        FROM 
                            Продажи
                        JOIN 
                            Клиенты ON Продажи.ID_Клиенты = Клиенты.ID
                        JOIN 
                            Сотрудники ON Продажи.ID_Сотрудники = Сотрудники.ID
                        JOIN 
                            Автомобили ON Продажи.ID_Автомобили = Автомобили.ID
                        JOIN 
                            Тип_модели ON Автомобили.ID_Тип_модели = Тип_модели.ID
                        JOIN 
                            Тип_марки ON Тип_модели.ID_Тип_марки = Тип_марки.ID
                        JOIN 
                            Тип_оплаты ON Продажи.ID_Тип_оплаты = Тип_оплаты.ID
                            Where Продажи.ID_Сотрудники = {self.id_users};
                        '''
        elif self.table_name == 'Продажи':
            sql_query = f'''SELECT 
                Продажи.ID AS ID_Продажи,
                Клиенты.ID AS ID_Клиента,
                Клиенты.Имя || ' ' || Клиенты.Фамилия AS Клиент,    
                Сотрудники.ID AS ID_Сотрудника,  
                Сотрудники.Имя || ' ' || Сотрудники.Фамилия AS Сотрудник,
                Автомобили.ID AS ID_Автомобиля,
                Автомобили.VIN AS VIN_Автомобиля,
                Тип_модели.Название AS Модель_Автомобиля,
                Тип_марки.Название AS Марка_Автомобиля,
                Продажи.Дата_продажи AS Дата_Продажи,
                Продажи.Цена AS Цена_Продажи,
                Тип_оплаты.Название AS Тип_Оплаты
            FROM 
                Продажи
            JOIN 
                Клиенты ON Продажи.ID_Клиенты = Клиенты.ID
            JOIN 
                Сотрудники ON Продажи.ID_Сотрудники = Сотрудники.ID
            JOIN 
                Автомобили ON Продажи.ID_Автомобили = Автомобили.ID
            JOIN 
                Тип_модели ON Автомобили.ID_Тип_модели = Тип_модели.ID
            JOIN 
                Тип_марки ON Тип_модели.ID_Тип_марки = Тип_марки.ID
            JOIN 
                Тип_оплаты ON Продажи.ID_Тип_оплаты = Тип_оплаты.ID
                ;
            '''

        else:
            sql_query = f'SELECT * FROM {self.table_name}'
        #sql_query = f'SELECT * FROM {self.table_name}'
        self.output_of_tables(sql_query, self.tableView)  # Выводим данные таблицы
        self.opr_cb_for_search()  # Обновляем список для поиска

    def search_column(self):
        try:
            value = self.le_name.text()  # Получаем текст из поля поиска
            column = self.cb_search.currentText()  # Получаем выбранную колонку

            if not value or not column or column == 'Выбирите поле':
                QMessageBox.information(self, 'Внимание', 'Выберите поле и введите \nданные!')
                return

            if self.table_name == 'Автомобили':
                column = f'Автомобили.{column}'
                sql_query = f"""SELECT 
                Автомобили.ID AS ID_Автомобиля,
                Тип_марки.Название AS Марка,
                Автомобили.ID_Тип_модели AS ID_Модели,
                Тип_модели.Название AS Модель,
                Автомобили.Год AS Год,
                Автомобили.VIN AS VIN,
                Автомобили.Цвет AS Цвет,
                Автомобили.Фото AS Фото,
                Автомобили.ID_Тип_двигателя AS ID_Двигателя,
                Тип_двигателя.Название AS Тип_двигателя,
                Автомобили.Объем AS Объем_двигателя,
                Автомобили.Мощность AS Мощность_двигателя,
                Автомобили.ID_Тип_трансмиссии AS ID_Трансмиссии,
                Тип_трансмиссии.Название AS Тип_трансмиссии,
                Автомобили.ID_Тип_привода AS ID_Привода,
                Тип_привода.Название AS Тип_привода,
                Автомобили.Масса AS Масса,
                Автомобили.Закуп AS Закупочная_цена,
                Автомобили.Цена AS Продажная_цена,
                Автомобили.Комплектация AS Комплектация,
                Автомобили.Статус AS Статус
            FROM 
                Автомобили
            JOIN 
                Тип_модели ON Автомобили.ID_Тип_модели = Тип_модели.ID
            JOIN 
                Тип_марки ON Тип_модели.ID_Тип_марки = Тип_марки.ID
            JOIN 
                Тип_двигателя ON Автомобили.ID_Тип_двигателя = Тип_двигателя.ID
            JOIN 
                Тип_трансмиссии ON Автомобили.ID_Тип_трансмиссии = Тип_трансмиссии.ID
            JOIN 
                Тип_привода ON Автомобили.ID_Тип_привода = Тип_привода.ID
            WHERE 
                {column} = '{value}' and Автомобили.Статус = 'Активный';
            """

            elif self.table_name == 'Продажи' and self.position == 'Сотрудник':
                column = f'Продажи.{column}'
                sql_query = f"""SELECT 
                Продажи.ID AS ID_Продажи,
                Продажи.ID_Клиенты AS ID_Клиента,
                Клиенты.Имя || ' ' || Клиенты.Фамилия AS Клиент,    
                Продажи.ID_Сотрудники AS ID_Сотрудника,  
                Сотрудники.Имя || ' ' || Сотрудники.Фамилия AS Сотрудник,
                Продажи.ID_Автомобили AS ID_Автомобиля,
                Автомобили.VIN AS VIN_Автомобиля,
                Тип_модели.Название AS Модель_Автомобиля,
                Тип_марки.Название AS Марка_Автомобиля,
                Продажи.Дата_продажи AS Дата_Продажи,
                Продажи.Цена AS Цена_Продажи,
                Тип_оплаты.Название AS Тип_Оплаты
            FROM 
                Продажи
            JOIN 
                Клиенты ON Продажи.ID_Клиенты = Клиенты.ID
            JOIN 
                Сотрудники ON Продажи.ID_Сотрудники = Сотрудники.ID
            JOIN 
                Автомобили ON Продажи.ID_Автомобили = Автомобили.ID
            JOIN 
                Тип_модели ON Автомобили.ID_Тип_модели = Тип_модели.ID
            JOIN 
                Тип_марки ON Тип_модели.ID_Тип_марки = Тип_марки.ID
            JOIN 
                Тип_оплаты ON Продажи.ID_Тип_оплаты = Тип_оплаты.ID
            WHERE {column} = '{value}' and Продажи.ID_Сотрудники = {self.id_users};"""

            elif self.table_name == 'Продажи':
                column = f'Продажи.{column}'
                sql_query = f"""SELECT 
                Продажи.ID AS ID_Продажи,
                Продажи.ID_Клиенты AS ID_Клиента,
                Клиенты.Имя || ' ' || Клиенты.Фамилия AS Клиент,    
                Продажи.ID_Сотрудники AS ID_Сотрудника,  
                Сотрудники.Имя || ' ' || Сотрудники.Фамилия AS Сотрудник,
                Продажи.ID_Автомобили AS ID_Автомобиля,
                Автомобили.VIN AS VIN_Автомобиля,
                Тип_модели.Название AS Модель_Автомобиля,
                Тип_марки.Название AS Марка_Автомобиля,
                Продажи.Дата_продажи AS Дата_Продажи,
                Продажи.Цена AS Цена_Продажи,
                Тип_оплаты.Название AS Тип_Оплаты
            FROM 
                Продажи
            JOIN 
                Клиенты ON Продажи.ID_Клиенты = Клиенты.ID
            JOIN 
                Сотрудники ON Продажи.ID_Сотрудники = Сотрудники.ID
            JOIN 
                Автомобили ON Продажи.ID_Автомобили = Автомобили.ID
            JOIN 
                Тип_модели ON Автомобили.ID_Тип_модели = Тип_модели.ID
            JOIN 
                Тип_марки ON Тип_модели.ID_Тип_марки = Тип_марки.ID
            JOIN 
                Тип_оплаты ON Продажи.ID_Тип_оплаты = Тип_оплаты.ID
            WHERE {column} = '{value}' ;"""

            else:
                sql_query = f"SELECT * FROM {self.table_name} WHERE {column} = '{value}';"

            self.output_of_tables(sql_query, self.tableView)  # Выводим результаты поиска

        except Exception as e:
            QMessageBox.information(self, 'Внимание', 'Заполните поле для поиска!')
            print(f"2 Ошибка при выполнении запроса: {e}")  # Обработка ошибок

    def opr_cb_for_search(self):
        self.cb_search.clear()  # Очищаем комбобокс
        print(self.datab.get_holder(self.table_name))
        data = self.datab.get_holder(self.table_name)
        if self.table_name == 'Автомобили':
            data.remove('Фото')
            data.remove('Статус')
            data.remove('Масса')
            data.remove('Закуп')
            data.remove('ID_Тип_привода')
            data.remove('ID_Тип_трансмиссии')

        print(data)
        for value in data:  # Заполняем новыми значениями
            self.cb_search.addItem(f"{value}")

    def output_of_tables(self, sql_query, tableView):
        try:
            self.datab.cursor.execute(sql_query)  # Выполнение запроса
            rows = self.datab.cursor.fetchall()  # Получение всех строк
            columns = [description[0] for description in self.datab.cursor.description]  # Получаем названия колонок

            # Создание модели данных для таблицы
            model = QStandardItemModel(len(rows), len(columns))
            model.setHorizontalHeaderLabels(columns)

            # Заполняем модель данными
            for row_idx, row_data in enumerate(rows):
                for col_idx, cell_data in enumerate(row_data):
                    item = QStandardItem(str(cell_data))
                    model.setItem(row_idx, col_idx, item)

            # Устанавливаем модель в таблицу
            tableView.setModel(model)
            tableView.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)

        except Exception as e:
            print(f"3 Ошибка при выполнении запроса: {e}")  # Обработка ошибок

    def delete_column(self):
        try:
            id_line, _, column = self.opr_cell_row_column(self.tableView)
            if not id_line or column != 0:
                QMessageBox.information(self, 'Внимание', 'Выберите строку для удаления!')
                return
            self.datab.query_delete_column(self.table_name, id_line)
            self.opr_sql_query()

        except Exception as e:
            print(f"4 Ошибка при выполнении запроса: {e}")  # Обработка ошибок

    def export_table_to_excel_with_header(self):
        directory = 'excel'
        filename = f'{self.table_name}_{date.today()}.xlsx'
        file = os.path.join(directory, filename)

        if not os.path.exists(directory):
            os.makedirs(directory)  # Создаем директорию, если её нет

        wb = Workbook()  # Создаём новый Excel-файл
        ws = wb.active

        model = self.tableView.model()  # Получаем модель данных таблицы
        if not model:
            return

        # Экспорт заголовков колонок
        header_labels = []
        for col in range(model.columnCount()):
            header_label = model.headerData(col, Qt.Orientation.Horizontal)
            header_labels.append(header_label)
            ws.cell(row=1, column=col + 1, value=str(header_label))

        # Экспорт строк таблицы
        for row in range(model.rowCount()):
            for col in range(model.columnCount()):
                item = model.index(row, col).data()
                if item is not None:
                    ws.cell(row=row + 2, column=col + 1, value=str(item))

        wb.save(file)  # Сохраняем Excel-файл
        QMessageBox.information(self, 'Внимание', 'Файл успешно создан!')

    def opr_cell_row_column(self, tableView ):
        selected_indexes = tableView.selectedIndexes()
        if not selected_indexes:
            return [None, None, None]  # или можно вернуть (None, None) в зависимости от вашего кода

        index = selected_indexes[0]

        cell = str(tableView.model().data(index))
        row = index.row()
        column = index.column()

        return [cell, row, column]

    def opr_click_tableview (self):
        try:
            image_path = self.opr_cell_row_column(self.tableView)[0]
            column = self.opr_cell_row_column(self.tableView)[2]

            ind = self.tableView.selectedIndexes()[0]
            id = str(self.tableView.model().data(ind))
            print(id)

            if self.table_name == 'Автомобили' and column == 7:
                image_window = ImageWindow(image_path)
                image_window.exec()

            elif self.table_name == 'sweets' and self.widget_add_ord.isVisible() and column == 0:
                name_unit = self.datab.query_opr_min_unit(id)[0]
                min_unit = self.datab.query_opr_min_unit(id)[1]
                print(name_unit,min_unit)
                self.widget_kol = widget_add_kol_vo(self.table_name, self)
                self.widget_kol.le_kol.setPlaceholderText(f' {name_unit}(минимум {min_unit})')
                self.widget_kol.show()
                self.widget_kol.but_confirm.clicked.connect(lambda: self.opr_kol_vo(id))

            elif self.table_name == 'sales' and column == 0:
                self.sales_or_delivers = ['sales', 'sale_items', 'sales_id']
                self.values_sales_delivers.append(int(id))
                self.open_widget_add_order()
                self.widget_add_ord.but_cancel_order.setVisible(False)
                self.view_sales_delivers(int(id))

            elif self.table_name == 'deliveries' and column == 0:
                self.sales_or_delivers = ['deliveries', 'delivery_items', 'deliveries_id']
                self.values_sales_delivers.append(int(id))
                self.open_widget_add_order()
                self.widget_add_ord.but_cancel_order.setVisible(False)
                self.view_sales_delivers(int(id))


        except Exception as e:
            print(f"5 Ошибка при выполнении запроса: {e}")

    def view_sales_delivers(self,id_sales_delivers ):
        try:
            # id_sales_delivers = self.values_sales_delivers[0]
            table_name = self.sales_or_delivers[0]

            if table_name == 'sales':
                sql_query = f'''SELECT sale_items.id, sweets.name, sweets.price, sale_items.quantity, unit.name as unit_name, manufacturers.name as manufacturers_name
                                        FROM sale_items JOIN sweets ON sale_items.sweets_id = sweets.id 
                                        JOIN unit on unit.id = sweets.unit_id 
                                        JOIN manufacturers ON manufacturers.id = sweets.manufacturers_id
                                        WHERE sale_items.sales_id = {id_sales_delivers};'''
            else:
                sql_query = f'''SELECT delivery_items.id, sweets.name, sweets.price, delivery_items.quantity, unit.name as unit_name, manufacturers.name as manufacturers_name
                                        FROM delivery_items JOIN sweets ON delivery_items.sweets_id = sweets.id 
                                        JOIN unit on unit.id = sweets.unit_id 
                                        JOIN manufacturers ON manufacturers.id = sweets.manufacturers_id
                                        WHERE delivery_items.deliveries_id = {id_sales_delivers};'''

            self.output_of_tables(sql_query, self.widget_add_ord.tableView_2)

            self.sum = f'{self.datab.query_sum_orders(self.sales_or_delivers[1], self.sales_or_delivers[2], id_sales_delivers)} Р'
            self.widget_add_ord.lb_sum_order_pr.setText(self.sum)
        except Exception as e:
            print(f"2 Ошибка при выполнении запроса: {e}")  # Обработка ошибок

    def opr_kol_vo(self, id):
        try:
            kol_vo = self.widget_kol.le_kol.text()
            min_unit = self.datab.query_opr_min_unit(id)[1]

            # Проверяем, что введенное значение является числом и больше или равно минимальной единице
            if kol_vo.isdigit() and int(kol_vo) >= min_unit:
                id_sales_delivers = self.values_sales_delivers[0]
                table_name = self.sales_or_delivers[1]
                column = self.sales_or_delivers[2]

                if not self.datab.query_repeat_check(id, id_sales_delivers, table_name, column, ):
                    self.widget_kol.le_kol.setText('')
                    self.widget_kol.le_kol.setPlaceholderText('Этот товар уже добавлен!')

                elif not self.datab.query_add_sweets_of_sales_delivers(id, id_sales_delivers, kol_vo ,table_name ):
                    self.widget_kol.le_kol.setText('')
                    self.widget_kol.le_kol.setPlaceholderText('Товара не хватает!')

                else:
                    self.widget_kol.close()
                    id_sales_delivers = self.values_sales_delivers[0]
                    self.view_sales_delivers(id_sales_delivers)

            else:
                self.widget_kol.le_kol.setText('')
                self.widget_kol.le_kol.setPlaceholderText('Неправильные данные!')
        except Exception as e:
            print(f"6 Ошибка при выполнении запроса: {e}")

    def show_about_info(self):
        QMessageBox.information(self, 'Информация о программе',
                                'Название: Система учета и анализа данных о покупках автомобилей.\n'
                                'Год создания: 2024')

    def closeEvent(self, event: QEvent):
        reply = QMessageBox(self)
        reply.setWindowTitle('Подтверждение')
        reply.setText('Вы уверены, что хотите закрыть приложение?')

        # Устанавливаем стандартные кнопки
        reply.setStandardButtons(QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

        # Меняем текст кнопок на русский язык
        reply.button(QMessageBox.StandardButton.Yes).setText('Да')
        reply.button(QMessageBox.StandardButton.No).setText('Нет')

        # Показываем диалог и проверяем, какая кнопка была нажата
        if reply.exec() == QMessageBox.StandardButton.Yes:
            event.accept()
        else:
            event.ignore()

    def close_all_windows(self):
        QApplication.closeAllWindows()


class widget_add_order(QDialog, Ui_Dialog_order):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setModal(True)
        #self.setWindowFlag(Qt.WindowType.WindowCloseButtonHint, False)
        self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowStaysOnTopHint)

class widget_report(QDialog, Ui_Dialog_report):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setModal(True)
        #self.setWindowFlag(Qt.WindowType.WindowCloseButtonHint, False)
        self.setWindowFlags(self.windowFlags() | Qt.WindowType.WindowStaysOnTopHint)

def main():
    app = QApplication([])
    window = LoginForm()
    window.show()
    app.exec()

if __name__ == "__main__":
    main()


